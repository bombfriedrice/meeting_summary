import streamlit as st
from openai import OpenAI
import os
from pydub import AudioSegment
from docx import Document

st.title("Meeting Note Generator")

# Set the API key and model name
MODEL = "gpt-4o"
client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))

uploaded_file = st.file_uploader("Upload your M4A recording", type=["m4a"])

if uploaded_file is not None:
    try:
        # Convert M4A to MP3 using pydub
        sound = AudioSegment.from_file(uploaded_file, format="m4a")
        sound.export("temp.mp3", format="mp3")
        audio_file_path = "temp.mp3"

        # Transcribe and translate the audio using Whisper model
        with open(audio_file_path, "rb") as audio_file:
            transcript = client.audio.transcriptions.create(
                model="whisper-1",
                file=audio_file
            )

        # Access the text field from the transcription result
        text = transcript.text

        st.write("Transcribed text:")
        st.write(text)

        # Download button for transcript
        st.download_button(
            label="Download Transcribed Text",
            data=text,
            file_name="meeting_transcript.txt",
            mime="text/plain",
        )

        # Summarize the transcription
        if st.button("Summarize"):
            summary_response = client.chat.completions.create(
                model=MODEL,
                messages=[
                    {
                        "role": "system",
                        "content": "You are generating a transcript summary. Create a summary of the provided transcription. Respond in Markdown."
                    },
                    {
                        "role": "user",
                        "content": f"The audio transcription is: {text}"
                    },
                ],
                temperature=0,
            )

            summary_text = summary_response.choices[0].message.content

            st.write("Meeting Summary:")
            st.write(summary_text)

            # Create a DOCX file with the transcription and summary
            doc = Document()
            doc.add_heading('Meeting Transcript', level=1)
            doc.add_paragraph(text)

            doc.add_heading('Meeting Summary', level=1)
            doc.add_paragraph(summary_text)

            # Save the DOCX file
            doc.save("meeting_notes.docx")

            # Download button for the DOCX file
            with open("meeting_notes.docx", "rb") as file:
                st.download_button(
                    label="Download Meeting Notes",
                    data=file,
                    file_name="meeting_notes.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )

    except Exception as e:
        st.write(f"Error during transcription: {e}")

    finally:
        # Clean up temporary MP3 file
        try:
            os.remove("temp.mp3")
        except FileNotFoundError:
            pass